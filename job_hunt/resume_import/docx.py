from __future__ import annotations

from pathlib import Path
from typing import Any


def _paragraph_text(paragraph) -> str:  # type: ignore[no-untyped-def]
    return "".join(node.text or "" for node in paragraph._p.iter() if node.tag.endswith("}t"))


def extract_docx(path: Path) -> tuple[str, list[str], dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install document support: pip install '.[documents]'") from exc
    try:
        document = Document(str(path))
    except Exception as exc:
        raise ValueError("DOCX is corrupted or unreadable") from exc
    lines: list[str] = []
    headings = 0
    list_items = 0
    for paragraph in document.paragraphs:
        text = _paragraph_text(paragraph).strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").casefold()
        if style.startswith("heading") or style.startswith("título"):
            level_text = "".join(character for character in style if character.isdigit())
            level = min(6, max(1, int(level_text or "2")))
            lines.append(f"{'#' * level} {text}")
            headings += 1
        elif "list" in style or "lista" in style:
            lines.append(f"- {text}")
            list_items += 1
        else:
            lines.append(text)
        lines.append("")
    table_count = 0
    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows or not any(any(cell for cell in row) for row in rows):
            continue
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        lines.extend(
            [
                "| " + " | ".join(normalized[0]) + " |",
                "| " + " | ".join(["---"] * width) + " |",
            ]
        )
        lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
        lines.append("")
        table_count += 1
    header_texts: list[str] = []
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = _paragraph_text(paragraph).strip()
            if text and text not in header_texts:
                header_texts.append(text)
    if header_texts:
        lines.extend(["## Cabeçalho", "", *header_texts, ""])
    metadata: dict[str, Any] = {
        "paragraphs": len(document.paragraphs),
        "headings": headings,
        "list_items": list_items,
        "tables": table_count,
        "headers": len(header_texts),
        "macros_executed": False,
    }
    return "\n".join(lines).strip(), [], metadata
