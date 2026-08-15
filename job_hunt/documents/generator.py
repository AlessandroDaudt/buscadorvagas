"""Generate factual, traceable Markdown and DOCX application documents."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from pydantic import Field

from job_hunt.domain.models import (
    JobAnalysisResult,
    MasterResume,
    ResumeExperience,
    StrictModel,
    UnifiedJob,
)
from job_hunt.normalization import normalize_match_text


class UnapprovedResumeError(ValueError):
    pass


class DocumentManifest(StrictModel):
    job_id: str
    job_url: str
    generated_at: datetime
    version: int = Field(ge=1)
    language: str
    model_used: str
    prompt_version: str
    master_hash: str
    files: list[str]
    changes: dict[str, object]


class GeneratedPackage(StrictModel):
    directory: Path
    resume_markdown: Path
    cover_letter_markdown: Path
    resume_docx: Path | None = None
    cover_letter_docx: Path | None = None
    resume_pdf: Path | None = None
    cover_letter_pdf: Path | None = None
    manifest: DocumentManifest


def _slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", normalize_match_text(value)).strip("-")[:80]


def _master_hash(master: MasterResume) -> str:
    content = json.dumps(master.model_dump(mode="json"), sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(content.encode()).hexdigest()


def _relevance(experience: ResumeExperience, job: UnifiedJob) -> int:
    haystack = normalize_match_text(
        " ".join(
            [
                experience.title,
                *experience.technologies,
                *experience.responsibilities,
                *experience.achievements,
            ]
        )
    )
    job_terms = set(normalize_match_text(f"{job.title} {job.description}").split())
    return sum(1 for term in job_terms if len(term) > 3 and term in haystack)


def _date_range(experience: ResumeExperience) -> str:
    if experience.start_date is None and experience.end_date is None:
        return ""
    start = experience.start_date.isoformat() if experience.start_date else ""
    end = experience.end_date.isoformat() if experience.end_date else "Present"
    return " - ".join(part for part in (start, end) if part)


def _tailor(master: MasterResume, job: UnifiedJob) -> tuple[MasterResume, dict[str, object]]:
    tailored = deepcopy(master)
    original_order = [experience.company for experience in tailored.experiences]
    tailored.experiences.sort(key=lambda experience: _relevance(experience, job), reverse=True)
    job_text = normalize_match_text(f"{job.title} {job.description}")
    tailored.skills.sort(
        key=lambda skill: (normalize_match_text(skill) in job_text, len(skill)), reverse=True
    )
    highlighted = [skill for skill in tailored.skills if normalize_match_text(skill) in job_text]
    changes: dict[str, object] = {
        "experience_order_before": original_order,
        "experience_order_after": [experience.company for experience in tailored.experiences],
        "highlighted_skills": highlighted,
        "facts_added": [],
        "facts_removed": [],
    }
    return tailored, changes


def render_resume_markdown(master: MasterResume) -> str:
    contact = master.contact
    contact_parts = [
        value
        for value in (
            contact.location,
            contact.email,
            contact.phone,
            contact.linkedin,
            contact.github,
        )
        if value
    ]
    lines = [f"# {contact.name}", " | ".join(contact_parts), "", "## Summary", "", master.summary]
    lines.extend(["", "## Experience"])
    for experience in master.experiences:
        date_range = _date_range(experience)
        suffix = f" | {date_range}" if date_range else ""
        lines.extend(["", f"### {experience.title} - {experience.company}{suffix}"])
        for bullet in [*experience.responsibilities, *experience.achievements]:
            lines.append(f"- {bullet}")
    lines.extend(["", "## Skills", "", ", ".join(master.skills)])
    lines.extend(["", "## Education"])
    for education in master.education:
        details = " - ".join(
            part
            for part in (education.qualification, education.institution, str(education.year or ""))
            if part
        )
        lines.append(f"- {details}")
    lines.extend(["", "## Certifications"])
    lines.extend(f"- {certification}" for certification in master.certifications)
    lines.extend(["", "## Languages"])
    lines.extend(f"- {language.language}: {language.level}" for language in master.languages)
    return "\n".join(lines).strip() + "\n"


def render_cover_letter(
    master: MasterResume,
    job: UnifiedJob,
    analysis: JobAnalysisResult,
    *,
    max_words: int = 350,
) -> str:
    english = master.language.lower().startswith("en")
    first = master.experiences[0] if master.experiences else None
    second = master.experiences[1] if len(master.experiences) > 1 else None
    gap = analysis.gaps[0] if analysis.gaps else None
    if english:
        paragraphs = [
            "Dear Hiring Team,",
            f"I am applying for the {job.title} position at {job.company}. {master.summary}",
        ]
        if first:
            paragraphs.append(
                f"My work as {first.title} at {first.company} includes "
                f"{first.responsibilities[0] if first.responsibilities else ', '.join(first.technologies[:5])}"
            )
        if second:
            paragraphs.append(
                f"Previously, as {second.title} at {second.company}, I developed further enterprise "
                f"support and troubleshooting experience across {', '.join(second.technologies[:6])}."
            )
        if gap:
            paragraphs.append(
                f"Although the analysis notes this non-critical gap - {gap} - my related enterprise "
                "security and troubleshooting background provides a practical foundation for learning it."
            )
        paragraphs.extend(
            [
                "I would welcome the opportunity to discuss how this experience can support your team.",
                f"Sincerely,\n{master.contact.name}",
            ]
        )
    else:
        paragraphs = [
            "Prezada equipe de recrutamento,",
            f"Apresento minha candidatura à posição de {job.title} na {job.company}. {master.summary}",
        ]
        if first:
            paragraphs.append(
                f"Minha atuação como {first.title} na {first.company} inclui "
                f"{first.responsibilities[0] if first.responsibilities else ', '.join(first.technologies[:5])}"
            )
        if second:
            paragraphs.append(
                f"Anteriormente, como {second.title} na {second.company}, ampliei minha experiência em "
                f"suporte enterprise e troubleshooting com {', '.join(second.technologies[:6])}."
            )
        if gap:
            paragraphs.append(
                f"Embora a análise aponte uma lacuna não crítica - {gap} - minha experiência relacionada "
                "em segurança e troubleshooting oferece uma base prática para desenvolvê-la."
            )
        paragraphs.extend(
            [
                "Fico à disposição para conversar sobre como essa experiência pode contribuir com a equipe.",
                f"Atenciosamente,\n{master.contact.name}",
            ]
        )
    content = "\n\n".join(paragraphs)
    words = content.split()
    if len(words) <= max_words:
        return content + "\n"
    return " ".join(words[:max_words]) + "\n"


def _write_docx(master: MasterResume, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Install document support: pip install '.[documents]'") from exc
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(1)
    styles["Title"].font.size = Pt(22)
    styles["Title"].paragraph_format.space_after = Pt(3)
    styles["Heading 1"].font.size = Pt(12)
    styles["Heading 1"].paragraph_format.space_before = Pt(5)
    styles["Heading 1"].paragraph_format.space_after = Pt(1)
    styles["Heading 2"].font.size = Pt(10.5)
    styles["Heading 2"].paragraph_format.space_before = Pt(3)
    styles["Heading 2"].paragraph_format.space_after = Pt(0)
    title = document.add_heading(master.contact.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = " | ".join(
        value
        for value in (
            master.contact.location,
            master.contact.email,
            master.contact.phone,
            master.contact.linkedin,
        )
        if value
    )
    paragraph = document.add_paragraph(contact)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("Summary", level=1)
    document.add_paragraph(master.summary)
    document.add_heading("Experience", level=1)
    for experience in master.experiences:
        document.add_heading(f"{experience.title} - {experience.company}", level=2)
        if _date_range(experience):
            document.add_paragraph(_date_range(experience))
        for bullet in [*experience.responsibilities, *experience.achievements]:
            document.add_paragraph(bullet, style="List Bullet")
    document.add_heading("Skills", level=1)
    document.add_paragraph(", ".join(master.skills))
    document.add_heading("Education", level=1)
    for education in master.education:
        document.add_paragraph(education.qualification, style="List Bullet")
    document.add_heading("Certifications", level=1)
    for certification in master.certifications:
        document.add_paragraph(certification, style="List Bullet")
    document.add_heading("Languages", level=1)
    for language in master.languages:
        document.add_paragraph(f"{language.language}: {language.level}", style="List Bullet")
    document.save(str(path))


def _write_cover_docx(content: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(11)
    for paragraph in content.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(str(path))


def export_pdf(docx_path: Path, *, timeout_seconds: int = 60) -> Path | None:
    executable = shutil.which("soffice")
    if not executable:
        return None
    subprocess.run(
        [
            executable,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        timeout=timeout_seconds,
    )
    output = docx_path.with_suffix(".pdf")
    return output if output.exists() else None


class DocumentGenerator:
    def __init__(self, master: MasterResume, *, output_root: Path = Path("output/doc")) -> None:
        if not master.approved:
            raise UnapprovedResumeError(
                "Master resume is not approved; review facts and set approved=true before generation"
            )
        self.master = master
        self.output_root = output_root

    def generate(
        self,
        job: UnifiedJob,
        analysis: JobAnalysisResult,
        *,
        create_docx: bool = True,
        create_pdf: bool = False,
        max_cover_words: int = 350,
        minimum_version: int | None = None,
    ) -> GeneratedPackage:
        tailored, changes = _tailor(self.master, job)
        directory = self.output_root / f"{_slug(job.company)}-{str(job.id)[:8]}"
        directory.mkdir(parents=True, exist_ok=True)
        existing = list(directory.glob("manifest-v*.json"))
        version = max(len(existing) + 1, minimum_version or 1)
        resume_md = directory / f"resume-v{version}.md"
        cover_md = directory / f"cover-letter-v{version}.md"
        resume_md.write_text(render_resume_markdown(tailored), encoding="utf-8")
        cover_content = render_cover_letter(tailored, job, analysis, max_words=max_cover_words)
        cover_md.write_text(cover_content, encoding="utf-8")
        files = [str(resume_md), str(cover_md)]
        resume_docx = cover_docx = resume_pdf = cover_pdf = None
        if create_docx:
            resume_docx = directory / f"resume-v{version}.docx"
            cover_docx = directory / f"cover-letter-v{version}.docx"
            _write_docx(tailored, resume_docx)
            _write_cover_docx(cover_content, cover_docx)
            files.extend([str(resume_docx), str(cover_docx)])
            if create_pdf:
                resume_pdf = export_pdf(resume_docx)
                cover_pdf = export_pdf(cover_docx)
                files.extend(str(path) for path in (resume_pdf, cover_pdf) if path)
        manifest = DocumentManifest(
            job_id=str(job.id),
            job_url=str(job.original_url),
            generated_at=datetime.now(timezone.utc),
            version=version,
            language=tailored.language,
            model_used="deterministic-factual-v1",
            prompt_version="documents-v1",
            master_hash=_master_hash(self.master),
            files=files,
            changes=changes,
        )
        manifest_path = directory / f"manifest-v{version}.json"
        manifest_path.write_text(manifest.model_dump_json(indent=2), encoding="utf-8")
        return GeneratedPackage(
            directory=directory,
            resume_markdown=resume_md,
            cover_letter_markdown=cover_md,
            resume_docx=resume_docx,
            cover_letter_docx=cover_docx,
            resume_pdf=resume_pdf,
            cover_letter_pdf=cover_pdf,
            manifest=manifest,
        )
