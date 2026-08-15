import json
from pathlib import Path

import pytest
from docx import Document
from pydantic import HttpUrl

from job_hunt.analysis.scoring import DeterministicScorer, consolidate_analysis
from job_hunt.configuration import load_candidate_profile, load_search_preferences
from job_hunt.documents.generator import DocumentGenerator, UnapprovedResumeError
from job_hunt.documents.importer import import_resume_candidate
from job_hunt.domain.models import MasterResume, UnifiedJob, WorkMode


def _master(approved=True):
    master = MasterResume.model_validate_json(
        Path("resume/master_resume.en.json").read_text(encoding="utf-8")
    )
    return master.model_copy(update={"approved": approved})


def _job():
    return UnifiedJob(
        source_name="fixture",
        original_url=HttpUrl("https://example.com/jobs/1"),
        company="Example Security",
        title="Senior Endpoint Security Engineer",
        description="Microsoft Defender for Endpoint, Entra ID, EDR, Windows, Linux and KQL",
        location="Remote Brazil",
        work_mode=WorkMode.REMOTE,
    )


def _analysis(job):
    return consolidate_analysis(
        DeterministicScorer(load_search_preferences(), load_candidate_profile()).score(job)
    )


def test_unapproved_master_resume_blocks_generation(tmp_path):
    with pytest.raises(UnapprovedResumeError, match="approved"):
        DocumentGenerator(_master(False), output_root=tmp_path)


def test_generate_traceable_markdown_and_docx_without_new_facts(tmp_path):
    job = _job()
    package = DocumentGenerator(_master(), output_root=tmp_path).generate(
        job, _analysis(job), create_docx=True
    )
    assert package.resume_markdown.exists()
    assert package.cover_letter_markdown.exists()
    assert package.resume_docx and package.resume_docx.exists()
    resume_text = package.resume_markdown.read_text(encoding="utf-8")
    assert "Microsoft Defender for Endpoint" in resume_text
    assert package.manifest.changes["facts_added"] == []
    manifest_file = package.directory / "manifest-v1.json"
    saved = json.loads(manifest_file.read_text(encoding="utf-8"))
    assert saved["job_id"] == str(job.id)
    assert len(saved["master_hash"]) == 64
    docx_text = "\n".join(p.text for p in Document(package.resume_docx).paragraphs)
    assert "Alessandro Luis Daudt" in docx_text
    assert "Dell Technologies" in docx_text


def test_generation_is_versioned_and_cover_letter_is_bounded(tmp_path):
    job = _job()
    generator = DocumentGenerator(_master(), output_root=tmp_path)
    first = generator.generate(job, _analysis(job), create_docx=False, max_cover_words=80)
    second = generator.generate(job, _analysis(job), create_docx=False, max_cover_words=80)
    assert first.manifest.version == 1
    assert second.manifest.version == 2
    cover = first.cover_letter_markdown.read_text(encoding="utf-8")
    assert len(cover.split()) <= 80
    assert job.company in cover
    assert job.title in cover


def test_generation_honors_persisted_minimum_version(tmp_path):
    job = _job()
    package = DocumentGenerator(_master(), output_root=tmp_path).generate(
        job,
        _analysis(job),
        create_docx=False,
        minimum_version=4,
    )
    assert package.manifest.version == 4
    assert (package.directory / "manifest-v4.json").exists()


def test_import_markdown_and_docx_are_unapproved_candidates(tmp_path):
    markdown = tmp_path / "resume.md"
    markdown.write_text("Approved facts still require review", encoding="utf-8")
    imported_md = import_resume_candidate(markdown)
    assert not imported_md.approved
    assert len(imported_md.source_hash) == 64

    docx_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Extracted DOCX fact")
    document.save(docx_path)
    imported_docx = import_resume_candidate(docx_path)
    assert "Extracted DOCX fact" in imported_docx.extracted_text
    unsupported = tmp_path / "resume.csv"
    unsupported.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError, match="supported"):
        import_resume_candidate(unsupported)
