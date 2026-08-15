import io
import zipfile

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from job_hunt.persistence.database import Database
from job_hunt.persistence.migration import upgrade_database
from job_hunt.persistence.models import ResumeVersionRecord
from job_hunt.resume_import.legacy_doc import LegacyDocConverterUnavailable
from job_hunt.resume_import.service import ResumeImportService
from job_hunt.resume_import.versions import ResumeVersionService


def test_text_and_docx_import_preserve_facts_and_structure(tmp_path):
    text = b"Maria Silva\n\nExperience\nAcme - Security Engineer - 2021\n\nEducation\nUniversity - 2019\n"
    result, safe_name, digest = ResumeImportService().import_bytes(
        text, original_filename="../Maria Resume.txt", content_type="text/plain"
    )
    assert safe_name == "Maria Resume.txt"
    assert "Acme - Security Engineer - 2021" in result.markdown
    assert result.source_format == "txt"
    assert len(digest) == 64
    assert {"experience", "education"}.issubset(result.detected_sections)

    document = Document()
    document.add_heading("Maria Silva", 1)
    document.add_heading("Experience", 2)
    document.add_paragraph("Acme - Security Engineer - 2021")
    document.add_paragraph("Python", style="List Bullet")
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "Python"
    path = tmp_path / "resume.docx"
    document.save(path)
    imported, _, _ = ResumeImportService().import_bytes(
        path.read_bytes(),
        original_filename="resume.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "## Experience" in imported.markdown
    assert "- Python" in imported.markdown
    assert imported.metadata["macros_executed"] is False


@pytest.mark.parametrize(
    ("name", "mime", "content", "message"),
    [
        ("resume.pdf", "application/pdf", b"not a PDF", "signature"),
        ("resume.txt", "text/plain", b"abc\x00def", "binary NUL"),
        ("resume.exe", "application/octet-stream", b"MZ...", "supported resume formats"),
    ],
)
def test_import_rejects_spoofed_or_binary_files(name, mime, content, message):
    with pytest.raises(ValueError, match=message):
        ResumeImportService().import_bytes(content, original_filename=name, content_type=mime)


def test_import_rejects_docx_macro_binary_and_pdf_without_text(tmp_path):
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w") as package:
        package.writestr("[Content_Types].xml", "<Types/>")
        package.writestr("word/document.xml", "<document/>")
        package.writestr("word/vbaProject.bin", b"macro")
    with pytest.raises(ValueError, match="macro-enabled"):
        ResumeImportService().import_bytes(
            archive.getvalue(),
            original_filename="resume.docx",
            content_type="application/zip",
        )

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    stream = io.BytesIO()
    writer.write(stream)
    with pytest.raises(ValueError, match="OCR local"):
        ResumeImportService().import_bytes(
            stream.getvalue(), original_filename="scan.pdf", content_type="application/pdf"
        )


def test_searchable_pdf_is_extracted_locally():
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 40 240 Td (Maria Security Engineer 2021) Tj ET")
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    page[NameObject("/Contents")] = content
    stream = io.BytesIO()
    writer.write(stream)

    result, safe_name, digest = ResumeImportService().import_bytes(
        stream.getvalue(), original_filename="searchable.pdf", content_type="application/pdf"
    )
    assert "Maria Security Engineer 2021" in result.markdown
    assert result.metadata["pages"] == 1
    assert safe_name == "searchable.pdf"
    assert len(digest) == 64


def test_password_protected_pdf_corrupt_docx_and_missing_doc_converter(tmp_path, monkeypatch):
    protected = PdfWriter()
    protected.add_blank_page(width=300, height=300)
    protected.encrypt("local-secret")
    stream = io.BytesIO()
    protected.write(stream)
    with pytest.raises(ValueError, match="protegido por senha"):
        ResumeImportService().import_bytes(
            stream.getvalue(), original_filename="protected.pdf", content_type="application/pdf"
        )

    with pytest.raises(ValueError, match="corrupted DOCX"):
        ResumeImportService().import_bytes(
            b"PK\x03\x04broken archive",
            original_filename="broken.docx",
            content_type="application/zip",
        )

    monkeypatch.setattr("job_hunt.resume_import.legacy_doc.shutil.which", lambda _name: None)
    ole_document = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + (b"\x00" * 64)
    with pytest.raises(LegacyDocConverterUnavailable, match="LibreOffice"):
        ResumeImportService().import_bytes(
            ole_document,
            original_filename="legacy.doc",
            content_type="application/msword",
        )


def test_resume_versions_require_approval_before_activation(tmp_path):
    url = f"sqlite:///{(tmp_path / 'versions.db').as_posix()}"
    upgrade_database(url)
    database = Database(url)
    text = b"# Maria\n\n## Experience\nAcme 2021\n\n## Education\nUniversity 2019\n"
    imported, safe_name, digest = ResumeImportService().import_bytes(
        text, original_filename="resume.md", content_type="text/markdown"
    )
    with database.session() as session:
        service = ResumeVersionService(session)
        first = service.create(imported, original_filename=safe_name, source_hash=digest)
        with pytest.raises(ValueError, match="approved"):
            service.activate(first)
        service.approve(first)
        service.activate(first)
        first_id = first.id
    with database.session() as session:
        current = session.get(ResumeVersionRecord, first_id)
        assert current is not None and current.approved and current.active
        second = ResumeVersionService(session).edit(current, current.markdown + "\nPython\n")
        assert second.previous_version_id == first_id
        assert not second.approved and not second.active
    database.dispose()
